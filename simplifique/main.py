import json
import time
import re
import sqlite3
import firebase_admin
import pyrebase
import requests
import wikipedia
from firebase_admin import auth as auth1
from firebase_admin import credentials, firestore
from kivy.core.text import LabelBase
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.screenmanager import ScreenManager, Screen
from kivymd.app import MDApp
from kivymd.uix.boxlayout import MDBoxLayout
from kivymd.uix.dialog import MDDialog
from kivymd.uix.list import ILeftBodyTouch, TwoLineAvatarIconListItem
from kivymd.uix.picker import MDDatePicker
from kivymd.uix.selectioncontrol import MDCheckbox
from kivy.uix.behaviors import ButtonBehavior
from kivy.uix.modalview import ModalView
from datetime import datetime
from kivymd.toast import toast
from docx import Document
from urllib.request import urlopen
from fileinput import close
from kivy.core.clipboard import Clipboard
from kivy.clock import Clock
from kivy.properties import ColorProperty, ListProperty, ObjectProperty, BooleanProperty, StringProperty, NumericProperty

#Libraries

cred = credentials.Certificate({
    "type": "service_account",
    "project_id": "simplifique-c881d",
    "private_key_id": "83f7f7868abc181d033a4d5e2c8ffe1c8af26688",
    "private_key": "-----BEGIN PRIVATE KEY-----\nMIIEvQIBADANBgkqhkiG9w0BAQEFAASCBKcwggSjAgEAAoIBAQDTi/op1fBF7Sx1\nv4+XkScBR42e31isc0olzMB/+ypswa4W7DmDywAxn30TASjWtCQeTdIBeiUeJstc\n7QGXku165kFnrUq6QdyOf6O/aGJibKHpuwY91NiFEO5/XkxwgmKbhUnjlWuNe8f1\nRB+o0Jiv833T8nfyiygtSaaDLgzlw0/lZ47NNfC9/kXprZrRFMDUhxlSB8wpHNQp\nbzQMLryYQx8/NfC4xr6NWI+v+aO32v65wGaXqpRpW7Tq81YBnXccND68EqKBJNTi\nfNIjmhbXj+DzFtgavK4oXPtHm4SSNq0q85jzc1e9oW6yS7gVpscrsXstOOdeunhm\nhTJ00qJnAgMBAAECggEAWTMirtVI1RNmtdeqQmywF7gGHUFr8HtEfp/RY6WSg/0+\n3OeKcOn+EO6BHKxWfgHDYQvLS1gnookVIv/Ethb8D+BbH85QPi1bRLLJZwIqyfmo\nZBe6UAQZsDItfoNSk/ZGgfE38MCmcygIboDlIJekajyvh9krfpfyvvXZQmL+iTqH\nZuqkqHLubfFaKEvL9jcrcYm1Z0ZLoGx4h9rLjc/0X8MMndkhuZYYn0CCn4i5fop8\nwxxsuACcIk1ug7ZiVJb6qMN2FSMEO7csghoCPHl0dP2RaaGpg2ILsDQhgGIWjRw3\nhdBj5sLXDdD/ZHlSburs5PjzxYPls9GRUfGQ8i5GMQKBgQDxtPwBob0uqM338LZN\niEcAMP5V70hGU8ryLXRRz3SCKmMxKJPeFsajVchfrNxexIQ/5MuWBSm9BO/g/qTN\nJDeFaQVPRIgTtT0xITNHcMUdT8rBwB32I7SwSkryIxsLUnaUAqA11WurUQk0cURb\nuFBZ3f6fDbQonl4oWwK2wR3+4wKBgQDgDmvHGLAve19A26fCXLE0Q87065RYLILh\nVTK2umNi4NPA8jWmxncI9F7eEBcEEph4aTlemmvBX2d3WNp5t/DyiTwszqrzTbeA\nikqrOqTxGJV24WvDehVI5zTVdqiKc0my+dJOUX5dbwV09I63nL77TWKXELfCemfw\nCbRL5v+BrQKBgGbLceG/x5VwdShdVyriKlAKhiBGA5blTApzCmVAtWwmWsktWLW7\nOf99HBqUiaREL3p885h52aZp0xr9MVmNbY9verKbksPO8JdUZ1qauzocFT8RVay4\nwr+22OjhxT6rc4K/GyPKAGB7tk53XXskiAewQfmi+lvL/n9rNVxEBV3BAoGAeEXF\nfT63dQWZAEvpJeB0D0ZHFhpPq3VZXHRLoOM07qMZiH18Z2YqB9iGBFZGxJzm09xI\nO4xRQ6Be/iXoQWaIJOmeL79Q7QJO+uVBZ+E3IWS89u/S1T/3pQbXya7Ekm2IplaM\nmhYM60Lpfvq4kb/GlUfZIJaMzgy/No8/BW+ewJECgYEArd8O6vFKY5y26zaMy6zN\nVi/3SdN2mXs96a0N6h+qhTcly53+P9Xpa/8xOZxypc26gf38JB/7VUaWX7FYDz20\n5GqleRXVSxmosaL9Hib9zGyaVlMlV/ytgIc150+C6SXJg5onwTZ+DKCuY7kBwIk6\nrx38QKyJd/+p5XbmTJcFPUo=\n-----END PRIVATE KEY-----\n",
    "client_email": "firebase-adminsdk-arl0m@simplifique-c881d.iam.gserviceaccount.com",
    "client_id": "110928610409474987418",
    "auth_uri": "https://accounts.google.com/o/oauth2/auth",
    "token_uri": "https://oauth2.googleapis.com/token",
    "auth_provider_x509_cert_url": "https://www.googleapis.com/oauth2/v1/certs",
    "client_x509_cert_url": "https://www.googleapis.com/robot/v1/metadata/x509/firebase-adminsdk-arl0m%40simplifique-c881d.iam.gserviceaccount.com"
    })
firebase_admin.initialize_app(cred)
class Accont:
    def __init__(self):
        firebaseConfig = {
            "apiKey": "AIzaSyCC_WAxsEv_GA99LBhucJ33_cWWy2aENCo",
            "authDomain": "simplifique-c881d.firebaseapp.com",
            "projectId": "simplifique-c881d",
            "databaseURL": "https://simplifique-c881d-default-rtdb.firebaseio.com",
            "storageBucket": "simplifique-c881d.appspot.com",
            "messagingSenderId": "413832219747",
            "appId": "1:413832219747:web:8108fd5f8557737180bc56",
            "measurementId": "G-2M4RF2PBPP"
        }
        firebase = pyrebase.initialize_app(firebaseConfig)
        self.auth = firebase.auth()

    def sign_in(self, email, password):
        try:
            self.auth.sign_in_with_email_and_password(email, password)
            return True

        except:
            return False
    def sing_up(self, email, password, name):
        try:
            auth1.create_user(email = email, email_verified = False, password = password, display_name = name, disabled = False)
            email = email
            user = auth1.get_user_by_email(email)
            user_id = user.uid
            display_name = user.display_name
            data = {'name': str(display_name), 'email': str(email)}
            db1 = firestore.client()
            db1.collection('users').document(user_id).set(data)
            return True
        except:
            return False
    def reset_password(self, email):
        try:
            self.auth.send_password_reset_email(email)
            return True

        except:
            return False
class Database:
    def __init__(self):
        self.con = sqlite3.connect('data.db')
        self.cursor = self.con.cursor()
        self.create_task_table()

    def create_task_table(self):
        """Create tasks table"""
        self.cursor.execute("CREATE TABLE IF NOT EXISTS tasks(id integer PRIMARY KEY AUTOINCREMENT, task varchar(50) NOT NULL, due_date varchar(50), completed BOOLEAN NOT NULL CHECK (completed IN (0, 1)))")
        self.con.commit()
        

    def create_task(self, task, due_date=None):
        """Create a task"""
        self.cursor.execute("INSERT INTO tasks(task, due_date, completed) VALUES(?, ?, ?)", (task, due_date, 0))
        self.con.commit()

        # GETTING THE LAST ENTERED ITEM SO WE CAN ADD IT TO THE TASK LIST
        created_task = self.cursor.execute("SELECT id, task, due_date FROM tasks WHERE task = ? and completed = 0", (task,)).fetchall()
        return created_task[-1]

    def get_tasks(self):
        """Get tasks"""
        uncomplete_tasks = self.cursor.execute("SELECT id, task, due_date FROM tasks WHERE completed = 0").fetchall()
        completed_tasks = self.cursor.execute("SELECT id, task, due_date FROM tasks WHERE completed = 1").fetchall()

        return completed_tasks, uncomplete_tasks

    

    def mark_task_as_complete(self, taskid):
        """Marking tasks as complete"""
        self.cursor.execute("UPDATE tasks SET completed=1 WHERE id=?", (taskid,))
        self.con.commit()

    def mark_task_as_incomplete(self, taskid):
        """Mark task as uncomplete"""
        self.cursor.execute("UPDATE tasks SET completed=0 WHERE id=?", (taskid,))
        self.con.commit()

        # return the text of the task
        task_text = self.cursor.execute("SELECT task FROM tasks WHERE id=?", (taskid,)).fetchall()
        return task_text[0][0]

    def delete_task(self, taskid):
        """Delete a task"""
        self.cursor.execute("DELETE FROM tasks WHERE id=?", (taskid,))
        self.con.commit()

    def close_db_connection(self):
        self.con.close()

#Screens

class Screen1(Screen):
    def on_pre_enter(self, *args):
        try:
            self.ids.l_email.text = ""
            self.ids.l_password.text = ""

        except:
            pass

    def try_sign_in(self):
        Clock.schedule_once(self.start_sign_in, .5)

    def start_sign_in(self, *args):
        email = self.ids.l_email.text
        password = self.ids.l_password.text

        if accont.sign_in(email, password):
            thing = MDApp.get_running_app()
            thing.root.current = "screen4"
            toast("Logged In Successfully!")
            user = auth1.get_user_by_email(email)
            user_id = user.uid
            name = user.display_name
            data = {
                'email': str(email),
                'password': str(password),
                'name': str(name),
                'uid': str(user_id)
                }
            json_object = json.dumps(data, indent = 4)
            
            # Writing to sample.json
            with open("cache.json", "w") as outfile:
                outfile.write(json_object)
            thing.press()
        else:
            if email and password:
                self.ids.l_password.text = ""
                self.ids.l_email.focus = True

            elif email and not password:
                self.ids.l_password.focus = True
                toast("This password isn't correct.")

            elif password and not email:
                self.ids.l_password.text = ""
                self.ids.l_email.focus = True
                toast("This email isn't in our database.")

            else:
                self.ids.l_email.focus = True
class Screen2(Screen):
    def on_pre_enter(self, *args):
        self.ids.s_email.text = ""
        self.ids.s_password.text = ""
        self.ids.s_name.text = ""

    def callback(self, *args):
        MDApp.get_running_app().root.current = "screen1"

    def validate_info(self):
        email = self.ids.s_email.text
        password = self.ids.s_password.text
        name = self.ids.s_name.text

        if password == password and "@" in email and ".com" in email and len(password) >= 6:
            if accont.sing_up(email, password, name):
                toast("Registration Created Successfully!")
                Clock.schedule_once(self.callback, 3)



            else:
                toast("Failed to create record.")

        else:
            if not email or "@" not in email or ".com" not in email:
                self.ids.s_email.focus = True

            elif len(password) < 6:
                self.ids.s_name.text = ""
                self.ids.s_password.text = ""
                self.ids.s_password.focus = True
class Screen3(Screen):
    def callback(self, *args):
        MDApp.get_running_app().root.current = "screen1"
    
    def on_pre_enter(self, *args):
        self.ids.r_email.text = ""

    def send_email_confirm(self):
        email = self.ids.r_email.text

        if accont.reset_password(email):
            toast("Email Succesfully Sended!")
        else:
            toast("Failed to send email.")

        Clock.schedule_once(self.callback, 3)
class Screen4(Screen):
    def callback(self, *args):
        MDApp.get_running_app().root.current = "screen1"
class Screen5(Screen):
    def callback(self, *args):
        MDApp.get_running_app().root.current = "screen4"
class Screen6(Screen):
    def callback(self, *args):
        MDApp.get_running_app().root.current = "screen4"
class Screen7(Screen):
    def callback(self, *args):
        MDApp.get_running_app().root.current = "screen4"
    def essay_helper(self, *args):
        w_name = self.ids.wiki_name.text
        language = self.ids.wiki_language.text
        title = self.ids.wiki_title.text
        wikipedia.set_lang(language)
        while True:
            try:
                wiki = wikipedia.page(title)
                break
            except:
                print("Project name invalid")
                title = input("Enter another project name: \n")
        text = wiki.content
        text = re.sub(r'==', '', text)
        text = re.sub(r'=', '', text)
        text = re.sub(r'\n', '\n    ', text)
        split = text.split('See also', 1)
        text = split[0]
        output_text = text   
        document = Document()
        paragraph = document.add_heading(title, 0)
        paragraph.alignment = 1
        paragraph = document.add_paragraph('    ' + text)
        paragraph = document.add_paragraph(w_name)
        paragraph.alignment = 2
        document.save(title + ".docx")
class Screen8(Screen):
    image_source = StringProperty('https://www.themealdb.com//images//media//meals//1520083578.jpg')
    def callback(self, *args):
        MDApp.get_running_app().root.current = "screen4"
    def  meal_text(self, *args):
        url = 'https://www.themealdb.com/api/json/v1/1/random.php'
        meal = requests.get(url).json()
        meals = meal["meals"]
        ingredients = []
        measures = []
        links = []
        meal_ingredients = ""
        global image_link
        for i in range(20):
            for ingredient in meals:
                if 1==1:
                    ingredients.append(ingredient[f"strIngredient{i+1}"])
                else:
                    break
            for measure in meals:
                if 1==1:
                    measures.append(measure[f"strMeasure{i+1}"])
                else:
                    break
        for j in range(20):
            try: 
                meal_ingredients = meal_ingredients + "• " + ingredients[j] + " - " + measures[j] + "\n"
            except:
                toast('Try again')
                pass
        for link in meals:
            links.append(link)
            image_link = link["strMealThumb"]
            yt_link = link["strYoutube"]
            meal_name = link['strMeal']
            meal_type = link['strCategory']
            meal_area = link['strArea']
            meal_instructions = link['strInstructions']
        try:
            meal_description = meal_name +'\n\n' + 'Category: ' + meal_type +'\n' + 'Area: ' + meal_area +'\n\n' + 'Instructions: \n\n' + meal_instructions +'\n'
            meal_ingredients_text = "Ingredients: \n\n" + meal_ingredients +  '\n' + "•  -" 
            meal_text = meal_description +'\n' + meal_ingredients_text.replace("•  -", "")
            self.ids.meal_text.text = meal_text
            self.image_source = image_link
        except:
            toast('Try Again')
            pass
        try:
            Clipboard.copy(yt_link)
        except:
            toast('Try Again')
            pass
        toast("Meal Found")
class Screen9(Screen):
    def callback(self, *args):
        MDApp.get_running_app().root.current = "screen4"
class Screen10(Screen):
    def callback(self, *args):
        MDApp.get_running_app().root.current = "screen4"
    def get_news(self, *args):
        lang = 'us'
        api_key = '283533136981441da324ba7c1b5d0cc5'
        url = f'https://newsapi.org/v2/top-headlines?country={lang}&apikey='+api_key
        news = requests.get(url).json()
        articles = news["articles"]
        my_articles = []
        my_news = ""
        for article in articles:
            my_articles.append(article["title"])
        for i in range(10):
            my_news = my_news + "• " + my_articles[i]+ "\n\n"
        text_news = my_news 
        self.ids.news_text.text = text_news   
        toast("Reload Succesfull")

class DialogContent(MDBoxLayout):
    """OPENS A DIALOG BOX THAT GETS THE TASK FROM THE USER"""
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        # set the date_text label to today's date when useer first opens dialog box
        self.ids.date_text.text = str(datetime.now().strftime('%A %d %B %Y'))


    def show_date_picker(self):
        """Opens the date picker"""
        date_dialog = MDDatePicker()
        date_dialog.bind(on_save=self.on_save)
        date_dialog.open()

    def on_save(self, instance, value, date_range):
        """This functions gets the date from the date picker and converts its it a
        more friendly form then changes the date label on the dialog to that"""

        date = value.strftime('%A %d %B %Y')
        self.ids.date_text.text = str(date)
class ListItemWithCheckbox(TwoLineAvatarIconListItem):
    '''Custom list item'''

    def __init__(self, pk=None, **kwargs):
        super().__init__(**kwargs)
        # state a pk which we shall use link the list items with the database primary keys
        self.pk = pk


    def mark(self, check, the_list_item):
        '''mark the task as complete or incomplete'''
        if check.active == True:
            the_list_item.text = '[s]'+the_list_item.text+'[/s]'
            db.mark_task_as_complete(the_list_item.pk)# here
        else:
            the_list_item.text = str(db.mark_task_as_incomplete(the_list_item.pk))# Here

    def delete_item(self, the_list_item):
        '''Delete the task'''
        self.parent.remove_widget(the_list_item)
        db.delete_task(the_list_item.pk)# Here
class LeftCheckbox(ILeftBodyTouch, MDCheckbox):
    '''Custom left container'''
class ListTile(ButtonBehavior, BoxLayout):
    title = StringProperty("")
    title_id = StringProperty("")
    icon = StringProperty("")
    subtitle = StringProperty("")
    extra = StringProperty("")
    amount = NumericProperty(0.0)
    expense = BooleanProperty(True)
    data = ObjectProperty(allowone=True)
    def __init__(self, **kw) -> None:
        super().__init__(**kw)
    def on_amount(self, inst, amount):
        amountx = self.ids.amount
        amountx.text.replace("-", "").replace("+", "")
        if self.expense:
            amountx.text = f"-{round(amount, 2)}"
        else:
            amountx.text = str(amount).strip()
        if amountx.text.startswith("+-"):
            amountx.text = "+"+amount.text[2:]
        if amountx.text.startswith("-+"):
            amountx.text = "-"+amount.text[2:]
    def on_expense(self, inst, amount):
        amount = self.ids.amount
        amount.text.replace("-", "").replace("+", "")
        if self.expense:
            amount.text = "-"+amount.text
        else:
            amount.text = "+"+amount.text
        if amount.text.startswith("+-"):
            amount.text = "+"+amount.text[2:]
        if amount.text.startswith("-+"):
            amount.text = "-"+amount.text[2:]
    def on_extra(self, inst, extra):
        self.ids.extra.text = f"${extra}"
class TileAction(ModalView):
    def __init__(self, **kw) -> None:
        super().__init__(**kw)
class AddNew(ModalView):
    expense = BooleanProperty(False)
    callback = ObjectProperty(print)
    def __init__(self, **kw) -> None:
        super().__init__(**kw)
    def confirm(self):
        self.dismiss()
        data = {
            'id': str(time.time()),
            'title': self.ids.t_title.text.strip(),
            'date': datetime.strftime(datetime.now(), "%Y-%m-%d, %H:%M:%S"),
            'amount': self.ids.t_amount.text.strip(),
            'initial-amount': '0.00',
            'icon': 'images/food.png', 
            'expense': self.expense,
            }
        self.callback(data)

#Main

class MainApp(MDApp, ScreenManager, BoxLayout, Screen10):
    global sm
    api_key = "60aae825eb1705b59a97532605cbae66"
    sm = ScreenManager()
    task_list_dialog = None
    dropdown = ObjectProperty
    def build(self):
        self.icon = "images/icon.png"
        self.title='Simplifique'
        self.theme_cls.theme_style = "Light"
        self.theme_cls.primary_palette = "Blue"
        sm.add_widget(Screen1(name='screen1')) # Login Page
        sm.add_widget(Screen2(name='screen2')) # SignUp Page
        sm.add_widget(Screen3(name='screen3')) # Reset Password
        sm.add_widget(Screen4(name='screen4')) # Home Screen
        sm.add_widget(Screen5(name='screen5')) # ToDo App
        sm.add_widget(Screen6(name='screen6')) # Weather App
        sm.add_widget(Screen7(name='screen7')) # Essay Helper
        sm.add_widget(Screen8(name='screen8')) # Financial App
        sm.add_widget(Screen9(name='screen9')) # Meal App
        sm.add_widget(Screen10(name='screen10')) # News App
        return sm
    def add_new(self, expense = True):
        an = AddNew()
        an.expense = expense
        an.open()
    def render(self, _):
        trans = [
            {
                'id': '6we',
                'title': 'Salary',
                'date': '6/10/2021',
                'amount': '12472',
                'initial-amount': '12,472',
                'icon': 'images/food.png', 
                'expense': False,
            },
            {
                'id': '6we',
                'title': 'Parents',
                'date': '10/10/2021',
                'amount': '125.00',
                'initial-amount': '12,472',
                'icon': 'images/food.png',
                'expense': False,
            },
            {
                'id': '6we',
                'title': 'Taxes',
                'date': '6/8/2021',
                'amount': '322.00',
                'initial-amount': '12,472',
                'icon': 'images/food.png', 
                'expense': True,
            },
            {
                'id': '6we',
                'title': 'Grand Parents',
                'date': '17/9/2022',
                'amount': '72.00',
                'initial-amount': '12,472',
                'icon': 'images/food.png',
                'expense': False,
            },
            {
                'id': '6we',
                'title': 'Grocceries',
                'date': '12/10/2022',
                'amount': '22.00',
                'initial-amount': '12,472',
                'icon': 'images/food.png',
                'expense': True
            }
        ]
        self.refresh_transaction(trans)
    def refresh_transaction(self, trans):
        screen9 = self.root.get_screen("screen9")
        grid = screen9.ids.gl_transactions
        grid.clear_widgets()
        for t in trans:
            tile = ListTile()
            tile.title_id = t['id']
            tile.title = t['title']
            tile.subtitle = t['date']
            tile.amount = t['amount']
            tile.extra = t['initial-amount']
            tile.expense = t['expense']
            tile.title = t['title']
            tile.icon = t['icon']
            tile.data = t
            tile.bind(on_release=self.tile_action)

            grid.add_widget(tile)
    def add_new(self, expense=True):
        an = AddNew()
        an.expense = expense
        an.callback = self.add_transaction
        an.open()
    def add_transaction(self, t):
        now = datetime.now()
        dt = datetime.strptime(t['date'], "%Y-%m-%d, %H:%M:%S")
        yr = now.year
        mnth = now.month
        day = now.day

        if yr == dt.year and mnth == dt.month:
            if day == dt.day:
                sub = "Today"
            elif dt.day == day -1:
                sub = "Yesterday"
        else:
            sub = t['date']

        tile = ListTile()
        tile.title_id = t['id']
        tile.title = t['title']
        tile.subtitle = sub
        tile.amount = t['amount']
        tile.extra = t['initial-amount']
        tile.expense = t['expense']
        tile.title = t['title']
        tile.icon = t['icon']
        tile.data = t
        tile.bind(on_release=self.tile_action)
        
        screen9_1 = self.root.get_screen("screen9")
        grid_1 = screen9_1.ids.gl_transactions
        grid_1.add_widget(tile)
    def tile_action(self, inst):
        ta = TileAction()
        ta.open()
    def show_task_dialog(self):
            if not self.task_list_dialog:
                self.task_list_dialog = MDDialog(
                    title="Create Task",
                    type="custom",
                    content_cls=DialogContent(),
                )

            self.task_list_dialog.open()
    def on_start(self):
        try:
            Clock.schedule_once(self.render, 1)
            completed_tasks, uncomplete_tasks = db.get_tasks()
            screen5 = self.root.get_screen("screen5")
            url = "http://ipinfo.io/json"
            response = urlopen(url)
            data = json.load(response)
            city = (data["city"])
            self.get_weather(city)

            if uncomplete_tasks != []:
                for task in uncomplete_tasks:
                    add_task = ListItemWithCheckbox(pk=task[0],text=task[1], secondary_text=task[2])
                    screen5.ids.container.add_widget(add_task)

            if completed_tasks != []:
                for task in completed_tasks:
                    add_task = ListItemWithCheckbox(pk=task[0],text='[s]'+task[1]+'[/s]', secondary_text=task[2])
                    add_task.ids.check.active = True
                    screen5.ids.container.add_widget(add_task)
        except requests.ConnectionError:
            print("No Internet Connection!")
            exit()
        except Exception as e :
            print(e)
            pass
    def close_dialog(self, *args):
        self.task_list_dialog.dismiss()
    def add_task(self, task, task_date):
        '''Add task to the list of tasks'''

        created_task = db.create_task(task.text, task_date)
        screen5 = self.root.get_screen("screen5")
        screen5.ids['container'].add_widget(ListItemWithCheckbox(pk=created_task[0], text='[b]'+created_task[1]+'[/b]', secondary_text=created_task[2]))# Here
        task.text = ''
    def get_weather(self, city_name):
        try:
            url = f"https://api.openweathermap.org/data/2.5/weather?q={city_name}&appid={self.api_key}"
            response = requests.get(url)
            x = response.json()
            screen6 = self.root.get_screen("screen6")
            if x["cod"] != "404":
                temperature = round(x["main"]["temp"]-273.15)
                humidity = x["main"]["humidity"]
                weather = x["weather"][0]["main"]
                id = str(x["weather"][0]["id"])
                wind_speed = round(x["wind"]["speed"]*18/5)
                location = x["name"] + ", " + x["sys"]["country"]
                screen6.ids.temperature.text = f"[b]{temperature}[/b]°"
                screen6.ids.weather.text = str(weather)
                screen6.ids.humidity.text = f"{humidity}%"
                screen6.ids.wind_speed.text = f"{wind_speed} km/h"
                screen6.ids.location.text = location
                if id == "800":
                    screen6.ids.weather_image.source = "images/sun.png"
                elif "200" <= id <= "232":
                    screen6.ids.weather_image.source = "images/storm.png"
                elif "300" <= id <= "321" and "500"<= id <= "531":
                    screen6.ids.weather_image.source = "images/rain.png"
                elif "600" <= id <= "622" :
                    screen6.ids.weather_image.source = "images/snow.png"
                elif "701" <= id <= "781":
                    screen6.ids.weather_image.source = "images/haze.png"
                elif "801" <= id <=  "804":
                    screen6.ids.weather_image.source = "images/clouds.png"
            else:
                toast("City Not Found")
        except requests.ConnectionError:
            toast("No Internet Connection!")
    def search_weather(self):
        screen6 = self.root.get_screen("screen6")
        city_name = screen6.ids.city_name.text
        if city_name != "":
            self.get_weather(city_name)   
    def press(self):
        with open('cache.json', 'r') as openfile:
            data = json.load(openfile)
        
        name1 = data["name"]
        email1 = data["email"]
        screen4 = self.root.get_screen("screen4")
        screen4.ids.home_text.text = f'Hello {name1}!'
        screen4.ids.drawer_text.text = name1
        screen4.ids.drawer_email.text = email1
        close
        
if __name__ == '__main__':
    LabelBase.register(name='Poppins', fn_regular='Poppins.ttf')
    LabelBase.register(name='Poppins-Bold', fn_regular='Poppins-Bold.ttf')
    accont = Accont()
    db = Database()
    app = MainApp()
    app.run()
